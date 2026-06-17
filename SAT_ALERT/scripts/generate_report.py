import os
from pathlib import Path
from docx import Document
from docx.shared import Pt

BASE = Path(__file__).resolve().parents[1]
README = BASE / 'README.txt'
REQUIREMENTS = BASE / 'requirements.txt'
SRC = BASE / 'src'
OUT_DIR = BASE / 'docs'
OUT_DIR.mkdir(exist_ok=True)
OUT_FILE = OUT_DIR / 'SAT_ALERT_Project_Report.docx'

FILES_TO_INCLUDE = [
    README,
    REQUIREMENTS,
]

# add key source files
for p in (SRC).glob('*.py'):
    FILES_TO_INCLUDE.append(p)

def add_preformatted(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


def read_text(path, max_lines=None):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            if max_lines:
                lines = []
                for _ in range(max_lines):
                    line = f.readline()
                    if not line:
                        break
                    lines.append(line)
                return ''.join(lines)
            return f.read()
    except Exception as e:
        return f'ERROR READING FILE: {e}'


doc = Document()

doc.add_heading('SAT_ALERT Project Report', level=1)
doc.add_paragraph('Auto-generated project report. Sections include README, requirements, and key source file excerpts.')

doc.add_page_break()

# README section
if README.exists():
    doc.add_heading('README', level=2)
    readme_text = read_text(README)
    for para in readme_text.split('\n\n'):
        p = doc.add_paragraph()
        add_preformatted(p, para + '\n')
    doc.add_page_break()

# Requirements
if REQUIREMENTS.exists():
    doc.add_heading('Requirements', level=2)
    req_text = read_text(REQUIREMENTS)
    p = doc.add_paragraph()
    add_preformatted(p, req_text)
    doc.add_page_break()

# Source files
src_files = sorted([p for p in FILES_TO_INCLUDE if p.exists() and p.suffix=='.py' and 'scripts' not in str(p)])
for path in src_files:
    rel = path.relative_to(BASE)
    doc.add_heading(f'File: {rel}', level=3)
    excerpt = read_text(path, max_lines=400)
    # add as preformatted paragraph preserving newlines
    for line in excerpt.splitlines():
        p = doc.add_paragraph()
        add_preformatted(p, line)
    doc.add_page_break()

# Footer
p = doc.add_paragraph()
add_preformatted(p, '\nReport generated from SAT_ALERT repo.\n')

# Save
try:
    doc.save(OUT_FILE)
    print(f'Report written to: {OUT_FILE}')
except Exception as e:
    # If file is locked (e.g., opened in Word/OneDrive), save to an alternate filename
    alt = OUT_DIR / 'SAT_ALERT_Project_Report_updated.docx'
    try:
        doc.save(alt)
        print(f'Report written to fallback file: {alt}')
    except Exception as e2:
        print('Failed to write DOCX:', e2)
